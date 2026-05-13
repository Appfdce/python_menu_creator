import pytest
from io import BytesIO
from docx import Document
from app.schemas.estimate_total import EstimateTotalRequest
from app.services.estimate_docx_generator import EstimateDocxGenerator
from app.services.estimate_perday_docx_generator import EstimatePerDayDocxGenerator

# Payload for testing food service layout
test_payload = {
    "event_id": "TEST-LAYOUT",
    "client": {"name": "Test Client"},
    "client_representative": {"name": "Rep Name"},
    "event": {
        "name": "Testing Layout",
        "guests": 70,
        "date_formatted": "Monday, June 15 2026",
        "end_date_formatted": "Monday, June 15 2026"
    },
    "meals": [
        {
            "show_date_header": True,
            "date_header": "June, Monday 15 2026",
            "category_name": "BREAKFAST",
            "category_precio_guest": "BREAKFAST @ $30.00 per Guest",
            "total_category_precio": "2100",  # Used by standard
            "total_category_precio_guest_por_dia": "2100",  # Used by perday
            "guest_count": "70",
            "date_day_name": "Monday"
        }
    ],
    "labor_services": [],
    "extras_events": [],
    "financials": {
        "total_food_service": "2100",
        "total_labor_cost": "0",
        "total_extras_events": "0",
        "total_estimate": "2100",
        "tax_rate": "0",
        "tax_name": "Tax",
        "service_charge_rate": "0",
        "credit_card_percent": "0"
    }
}

def test_food_service_layout_standard():
    req = EstimateTotalRequest(**test_payload)
    generator = EstimateDocxGenerator()
    
    docx_bytes = generator.generate_docx(req)
    doc = Document(docx_bytes)
    
    # Collect text and font properties in the Food Service section
    # The food service section is after the "Food Service" text
    food_section_started = False
    collected_paragraphs = []
    
    for p in doc.paragraphs:
        if "Food Service" in p.text:
            food_section_started = True
            continue
        
        # We assume food service section ends when we hit Cost of Balance or other section
        if food_section_started and ("Cost of Balance" in p.text or "Labor Service Fees" in p.text):
            break
            
        if food_section_started and p.text.strip():
            collected_paragraphs.append(p)
            
    texts = [p.text for p in collected_paragraphs]
    
    # Assertions for Standard Generator:
    # 1. The item "BREAKFAST @ $30.00 per Guest" should exist.
    # 2. The daily total "$ 2,100.00" should exist.
    # 3. The daily total MUST appear AFTER the breakfast line item.
    
    item_index = -1
    total_index = -1
    
    for idx, text in enumerate(texts):
        if "BREAKFAST @ $30.00 per Guest" in text:
            item_index = idx
            # Verify that the price in the item row is NOT bold
            # For item paragraph: usually it has "BREAKFAST @ $30.00 per Guest \t $ 2,100.00"
            # We need to find the run containing "$ 2,100.00"
            item_para = collected_paragraphs[idx]
            price_run_found = False
            for r in item_para.runs:
                if "$ 2,100.00" in r.text:
                    price_run_found = True
                    assert not r.bold, f"Expected individual price run NOT to be bold, but it is. Run text: '{r.text}'"
            assert price_run_found, f"Price run not found in item row! Paragraph text: '{item_para.text}'"
            
        # The total line should be only the total, possibly with leading tabs
        elif text.strip() == "$ 2,100.00":
            total_index = idx
            # Verify total is bold in at least one run
            total_para = collected_paragraphs[idx]
            total_run_bold = False
            for r in total_para.runs:
                if "$ 2,100.00" in r.text and r.bold:
                    total_run_bold = True
            assert total_run_bold, f"Expected final total run to be bold, but it's not. Runs: {[r.text + ' (bold=' + str(r.bold) + ')' for r in total_para.runs]}"
            
    assert item_index != -1, f"Could not find item row in {texts}"
    assert total_index != -1, f"Could not find total row in {texts}"
    assert total_index > item_index, f"Expected daily total to be AFTER the breakfast item. Indices: item={item_index}, total={total_index}"


def test_food_service_layout_perday():
    req = EstimateTotalRequest(**test_payload)
    generator = EstimatePerDayDocxGenerator()
    
    docx_bytes = generator.generate_docx(req)
    doc = Document(docx_bytes)
    
    food_section_started = False
    collected_paragraphs = []
    
    for p in doc.paragraphs:
        if "Food Service" in p.text:
            food_section_started = True
            continue
        
        if food_section_started and ("Cost of Balance" in p.text or "Labor Service Fees" in p.text):
            break
            
        if food_section_started and p.text.strip():
            collected_paragraphs.append(p)
            
    texts = [p.text for p in collected_paragraphs]
    
    item_index = -1
    total_index = -1
    
    for idx, text in enumerate(texts):
        if "BREAKFAST @ $30.00 per Guest" in text:
            item_index = idx
            item_para = collected_paragraphs[idx]
            price_run_found = False
            for r in item_para.runs:
                if "$ 2,100.00" in r.text:
                    price_run_found = True
                    assert not r.bold, f"Expected individual price run NOT to be bold, but it is. Run text: '{r.text}'"
            assert price_run_found, f"Price run not found in item row! Paragraph text: '{item_para.text}'"
            
        elif text.strip() == "$ 2,100.00":
            total_index = idx
            # Verify total is bold in at least one run
            total_para = collected_paragraphs[idx]
            total_run_bold = False
            for r in total_para.runs:
                if "$ 2,100.00" in r.text and r.bold:
                    total_run_bold = True
            assert total_run_bold, f"Expected final total run to be bold, but it's not. Runs: {[r.text + ' (bold=' + str(r.bold) + ')' for r in total_para.runs]}"
            
    assert item_index != -1, f"Could not find item row in {texts}"
    assert total_index != -1, f"Could not find total row in {texts}"
    assert total_index > item_index, f"Expected daily total to be AFTER the breakfast item. Indices: item={item_index}, total={total_index}"
