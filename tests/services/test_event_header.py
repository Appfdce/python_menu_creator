import pytest
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from app.schemas.estimate_total import EstimateTotalRequest
from app.services.estimate_docx_generator import EstimateDocxGenerator
from app.services.estimate_perday_docx_generator import EstimatePerDayDocxGenerator

event_test_payload = {
    "event_id": "TEST-EVENT-HEADER",
    "client": {"name": "Test Client"},
    "client_representative": {"name": "Rep Name"},
    "event": {
        "name": "Special Purple Event Gala",
        "address": "1234 Purple Avenue, Lavender City, FL 98765",
        "guests": 10,
        "date_formatted": "Tuesday, Oct 27 2026",
        "end_date_formatted": "Friday, Oct 30 2026"
    },
    "meals": [
        {
            "date_header": "October, Tuesday 27 2026",
            "category_name": "Breakfast",
            "total_category_precio": "100",
            "total_category_precio_guest_por_dia": "100",
            "category_precio_guest": "Breakfast",
            "guest_count": 10,
            "provide_by_client": False,
            "subcategory_1_items": []
        }
    ],
    "labor_services": [],
    "extras_events": [],
    "financials": {
        "tax_rate": "0,000 %",
        "tax_name": "Tax Exempt",
        "service_charge_rate": "10,00 %",
        "credit_card_percent": "0,00 %"
    }
}

def check_event_header_assertions(docx_bytes):
    doc = Document(docx_bytes)
    
    # Find the index of the paragraph that says "MENUS"
    menus_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == "MENUS":
            menus_idx = idx
            break
            
    assert menus_idx != -1, "Could not find 'MENUS' header in the document"
    
    # The event name and address should be just above "MENUS"
    # Wait, there might be empty lines or other text before MENUS, 
    # but at the start of the dynamic content we expect:
    # Event Name paragraph
    # Event Address paragraph
    # So menus_idx - 2 and menus_idx - 1 should be our targets.
    
    assert menus_idx >= 2, f"'MENUS' index {menus_idx} is too small to contain event header before it"
    
    p_name = doc.paragraphs[menus_idx - 2]
    p_address = doc.paragraphs[menus_idx - 1]
    
    assert p_name.text.strip() == "Special Purple Event Gala", f"Expected Event Name but got '{p_name.text}'"
    assert p_address.text.strip() == "1234 Purple Avenue, Lavender City, FL 98765", f"Expected Event Address but got '{p_address.text}'"
    
    # Verify centering
    assert p_name.alignment == WD_ALIGN_PARAGRAPH.CENTER, "Event name is not centered"
    assert p_address.alignment == WD_ALIGN_PARAGRAPH.CENTER, "Event address is not centered"
    
    # Verify Event Name Run properties
    assert len(p_name.runs) > 0
    run_name = p_name.runs[0]
    
    # Color should be primary (Wine 0x612d4b)
    # R=0x61 (97), G=0x2d (45), B=0x4b (75)
    assert run_name.font.color.rgb == RGBColor(0x61, 0x2d, 0x4b), "Event name color is not the primary purple/wine"
    
    # Font size of address should be smaller than 10pt or at least smaller than the event name.
    assert len(p_address.runs) > 0
    run_address = p_address.runs[0]
    
    # If standard text is 10pt, name could be larger (e.g. Pt(12) or Pt(14)), 
    # and address could be slightly smaller than name, e.g. Pt(9) or Pt(10).
    # The assertion enforces that run_address font size is less than run_name font size.
    assert run_address.font.size < run_name.font.size, "Event address font size is not smaller than event name"

def test_event_header_in_docx_generator():
    req = EstimateTotalRequest(**event_test_payload)
    generator = EstimateDocxGenerator()
    docx_bytes = generator.generate_docx(req)
    check_event_header_assertions(docx_bytes)

def test_event_header_in_perday_docx_generator():
    req = EstimateTotalRequest(**event_test_payload)
    generator = EstimatePerDayDocxGenerator()
    docx_bytes = generator.generate_docx(req)
    check_event_header_assertions(docx_bytes)
