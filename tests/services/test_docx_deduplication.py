import pytest
from io import BytesIO
from docx import Document
from app.schemas.estimate_total import EstimateTotalRequest
from app.services.estimate_docx_generator import EstimateDocxGenerator
from app.services.estimate_perday_docx_generator import EstimatePerDayDocxGenerator

# Sample payload with subtle variations that SHOULD be considered duplicates but might fail simple string checks
test_payload = {
    "event_id": "TEST-DEDUPLICATION",
    "client": {"name": "Test Client"},
    "client_representative": {"name": "Rep Name"},
    "event": {
        "name": "Testing Duplicates",
        "guests": 100,
        "date_formatted": "Tuesday, Oct 27 2026",
        "end_date_formatted": "Friday, Oct 30 2026"
    },
    "meals": [],
    "labor_services": [
        {
            "show_date_header": True,
            "date_header": "October, Tuesday 27 2026 ",  # Trailing space
            "hours": "6,00",
            "name": "1 Bartender",
            "total": "330"  # Different total string
        },
        {
            "show_date_header": False,
            "date_header": "October, Tuesday 27 2026",
            "hours": "6.00 ",  # Trailing space and dot instead of comma
            "name": "1 Bartender ",  # Trailing space
            "total": "$ 330.00"  # Different total string
        }
    ],
    "extras_events": [],
    "financials": {
        "total_food_service": "0",
        "total_labor_cost": "330",
        "total_extras_events": "0",
        "total_estimate": "330",
        "tax_rate": "0",
        "tax_name": "Tax",
        "service_charge_rate": "0",
        "credit_card_percent": "0"
    }
}

def test_labor_deduplication_docx_generator():
    req = EstimateTotalRequest(**test_payload)
    generator = EstimateDocxGenerator()
    
    docx_bytes = generator.generate_docx(req)
    doc = Document(docx_bytes)
    
    # Count occurrences of "Bartender" in the document text
    full_text = "\n".join([p.text for p in doc.paragraphs])
    bartender_count = full_text.count("Bartender")
    
    # Since they are duplicates, we only expect 1 bartender text block
    assert bartender_count == 1, f"Expected 1 bartender, but found {bartender_count}. Full text: \n{full_text}"

def test_labor_deduplication_perday_docx_generator():
    req = EstimateTotalRequest(**test_payload)
    generator = EstimatePerDayDocxGenerator()
    
    docx_bytes = generator.generate_docx(req)
    doc = Document(docx_bytes)
    
    # Count occurrences of "Bartender" in the document text
    full_text = "\n".join([p.text for p in doc.paragraphs])
    bartender_count = full_text.count("Bartender")
    
    # Since they are duplicates, we only expect 1 bartender text block
    assert bartender_count == 1, f"Expected 1 bartender, but found {bartender_count}. Full text: \n{full_text}"

def test_labor_date_ordering_and_headers():
    payload = {
        "event_id": "TEST-ORDERING",
        "client": {"name": "Test Client"},
        "client_representative": {"name": "Rep Name"},
        "event": {
            "name": "Testing Ordering",
            "guests": 100,
            "date_formatted": "Tuesday, Oct 27 2026",
            "end_date_formatted": "Friday, Oct 30 2026"
        },
        "meals": [],
        "labor_services": [
            {
                "show_date_header": True,
                "date_header": "June, Wednesday 17 2026", # Out of order, should be last chronologically
                "hours": "10,00",
                "name": "1 Event Manager",
                "total": "500"
            },
            {
                "show_date_header": False, # AppSheet gave False, but it IS the first/only of June 15!
                "date_header": "June, Monday 15 2026",
                "hours": "6,00",
                "name": "1 Bartender",
                "total": "330"
            },
            {
                "show_date_header": False,
                "date_header": "June, Tuesday 16 2026",
                "hours": "6,00",
                "name": "1 Bartender",
                "total": "330"
            }
        ],
        "extras_events": [],
        "financials": {
            "total_food_service": "0",
            "total_labor_cost": "1160",
            "total_extras_events": "0",
            "total_estimate": "1160",
            "tax_rate": "0",
            "tax_name": "Tax",
            "service_charge_rate": "0",
            "credit_card_percent": "0"
        }
    }
    req = EstimateTotalRequest(**payload)
    generator = EstimateDocxGenerator()
    
    docx_bytes = generator.generate_docx(req)
    doc = Document(docx_bytes)
    
    full_text = "\n".join([p.text for p in doc.paragraphs])
    
    # Assert headers for all three distinct dates are present in the output
    assert "June, Monday 15 2026" in full_text, f"Missing June 15 header! Text:\n{full_text}"
    assert "June, Tuesday 16 2026" in full_text, f"Missing June 16 header! Text:\n{full_text}"
    assert "June, Wednesday 17 2026" in full_text, f"Missing June 17 header! Text:\n{full_text}"
    
    # Assert chronological order: June 15 should appear before June 16, which is before June 17
    idx_15 = full_text.index("June, Monday 15 2026")
    idx_16 = full_text.index("June, Tuesday 16 2026")
    idx_17 = full_text.index("June, Wednesday 17 2026")
    
    assert idx_15 < idx_16, "June 15 should appear BEFORE June 16 in the document!"
    assert idx_16 < idx_17, "June 16 should appear BEFORE June 17 in the document!"
