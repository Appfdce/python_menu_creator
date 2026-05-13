import pytest
from io import BytesIO
from docx import Document
from app.schemas.estimate_total import EstimateTotalRequest
from app.services.estimate_docx_generator import EstimateDocxGenerator
from app.services.estimate_perday_docx_generator import EstimatePerDayDocxGenerator

# Construct a robust payload for financial testing
financial_test_payload = {
    "event_id": "TEST-FINANCIALS",
    "client": {"name": "Test Client"},
    "client_representative": {"name": "Rep Name"},
    "event": {
        "name": "Testing Financial Rules",
        "guests": 10,
        "date_formatted": "Tuesday, Oct 27 2026",
        "end_date_formatted": "Friday, Oct 30 2026"
    },
    # 1 meal costing $100 total
    "meals": [
        {
            "date_header": "October, Tuesday 27 2026",
            "category_name": "Breakfast",
            "total_category_precio": "100",
            "total_category_precio_guest_por_dia": "100",
            "category_precio_guest": "Breakfast",
            "guest_count": 10,
            "provide_by_client": False,
            "subcategory_1_items": [{"name": "Eggs", "total_item_price": "100", "guest_count": 10, "item_precio_guest": "Eggs"}]
        }
    ],
    # Labor costing $50 total
    "labor_services": [
        {
            "date_header": "October, Tuesday 27 2026",
            "hours": "5,00",
            "name": "Bartender",
            "total": "50"
        }
    ],
    # Rentals costing $1000, no sales extra
    "extras_events": [
        {
            "date_header": "October, Tuesday 27 2026",
            "is_rental": True,
            "is_sales": False,
            "name": "Tent Rental",
            "name_rental": "Tent Rental",
            "name_sales": "",
            "total": "1000",
            "provide_by_client": False
        }
    ],
    "financials": {
        "tax_rate": "0,000 %",
        "tax_name": "Tax Exempt",
        "service_charge_rate": "10,00 %",
        "credit_card_percent": "0,00 %"
    }
}

def check_financial_assertions(docx_bytes):
    doc = Document(docx_bytes)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    
    # 1. Zero-value items should NOT show
    # Extras Services (Sales) has 0.00 total, should not show!
    assert "Extras Services (Sales)" not in full_text, "Extras Services (Sales) with $0.00 should be hidden!"
    
    # Tax has 0.00 total (0.00%), should not show!
    assert "Tax Exempt" not in full_text, "Tax with 0.00% ($0.00) should be hidden!"
    
    # 2. Event Subtotal (Pre-Tax) must not appear in text at all
    assert "Event Subtotal (Pre-Tax)" not in full_text, "Event Subtotal (Pre-Tax) must be hidden from view!"

    # 3. Extras Services (Rentals) should be normal currency (no parentheses or asterisk)
    # It has value 1000.00
    assert "$ 1,000.00" in full_text, "Rentals value of $1,000.00 should be visible."
    # It must not have parentheses with asterisk
    assert "($ 1,000.00*)" not in full_text, "Rentals must not be marked as informative with parentheses/asterisks!"

    # 4. Mathematical verification
    # Food: 100.00
    # Labor: 50.00
    # Subtotal 1 (Pre-Tax) = 150.00
    # Subtotal 2 (After tax) = 150.00 + 0% = 150.00
    # Service Charge = (100 + 50) * 10% = 15.00
    # Rentals = 1000.00
    # Total Estimated Amount should be 150.00 + 1000.00 + 15.00 = 1165.00
    assert "Total Estimated Amount" in full_text
    
    # Find the line containing "Total Estimated Amount" and verify it contains 1,165.00
    found_amount = False
    for p in doc.paragraphs:
        if "Total Estimated Amount" in p.text:
            assert "1,165.00" in p.text, f"Expected total 1,165.00 but got {p.text}"
            found_amount = True
            break
    assert found_amount, "Could not find 'Total Estimated Amount' paragraph!"

def test_financial_rules_estimate_docx_generator():
    req = EstimateTotalRequest(**financial_test_payload)
    generator = EstimateDocxGenerator()
    docx_bytes = generator.generate_docx(req)
    check_financial_assertions(docx_bytes)

def test_financial_rules_estimate_perday_docx_generator():
    req = EstimateTotalRequest(**financial_test_payload)
    generator = EstimatePerDayDocxGenerator()
    docx_bytes = generator.generate_docx(req)
    check_financial_assertions(docx_bytes)
