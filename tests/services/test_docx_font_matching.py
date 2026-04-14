
import pytest
from io import BytesIO
from docx import Document
from app.services.estimate_docx_generator import EstimateDocxGenerator
from app.schemas.estimate_total import EstimateTotalRequest

@pytest.fixture
def sample_request():
    return EstimateTotalRequest(**{
        "event_id": "test",
        "client": {"name": "Test Client", "address": "Addr", "email": "e@e.com"},
        "client_representative": {"name": "Rep", "email": "r@r.com", "formatted_phone": "123"},
        "event": {
            "name": "Event", 
            "address": "Addr", 
            "code": "C", 
            "date_formatted": "Wed", 
            "end_date_formatted": "Thu", 
            "guests": 10
        },
        "meals": [],
        "labor_services": [],
        "extras_events": [],
        "financials": {
            "total_food_service": "$0",
            "total_labor_cost": "$0",
            "total_extras_events": "$0",
            "tax_name": "Tax",
            "tax_rate": "0%",
            "total_tax": "$0",
            "service_charge_rate": "0%",
            "total_service_charge": "$0",
            "total_estimate": "$0"
        }
    })

def test_font_is_open_sans(sample_request):
    generator = EstimateDocxGenerator()
    docx_stream = generator.generate_docx(sample_request)
    doc = Document(docx_stream)
    
    # Check if we have dynamic paragraphs and if they use Open Sans
    # The generator adds "MENUS" as the first dynamic paragraph
    found_dynamic = False
    for p in doc.paragraphs:
        if "MENUS" in p.text:
            found_dynamic = True
            for run in p.runs:
                if "MENUS" in run.text:
                    # Specifically check the rFonts element which is more reliable
                    rFonts = run._element.rPr.rFonts if run._element.rPr is not None else None
                    if rFonts is not None:
                        # We want it to be 'Open Sans' in ascii/hAnsi
                        assert rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii') == "Open Sans"
                    else:
                        # Fallback check
                        assert run.font.name == "Open Sans"
    
    assert found_dynamic, "Dynamic content marker not found or paragraphs not added"

def test_title_color_matches_html(sample_request):
    generator = EstimateDocxGenerator()
    docx_stream = generator.generate_docx(sample_request)
    doc = Document(docx_stream)
    
    for p in doc.paragraphs:
        if "MENUS" in p.text:
            for run in p.runs:
                if "MENUS" in run.text:
                    # HTML color #612d4b -> (97, 45, 75)
                    color = run.font.color.rgb
                    assert color == (0x61, 0x2d, 0x4b)

def test_placeholder_font_is_open_sans(sample_request):
    generator = EstimateDocxGenerator()
    docx_stream = generator.generate_docx(sample_request)
    doc = Document(docx_stream)
    
    # Check "Test Client" which replaces {{CLIENT_NAME}}
    found = False
    for p in doc.paragraphs:
        if "Test Client" in p.text:
            found = True
            for run in p.runs:
                if "Test Client" in run.text:
                    # In current implementation, this will likely fail or return None/Default
                    # because it's just doing run.text replacement on existing runs.
                    assert run.font.name == "Open Sans"
    
    assert found, "Placeholder value not found"
