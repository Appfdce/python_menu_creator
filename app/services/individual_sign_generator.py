import os
import logging
from io import BytesIO
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from app.schemas.individual_menu import IndividualSignRequest

logger = logging.getLogger(__name__)

LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "static", "logo_fdce.png")

class ItemData:
    def __init__(self, name, description, diet_options):
        self.name = name
        self.description = description
        self.diet_options = diet_options


def format_cell(cell, item):
    """Formats a table cell as a label."""
    # Set cell dimensions (approx 8cm x 5cm)
    # Note: python-docx doesn't always strictly enforce cell width if table is auto-layout
    # but we will try.
    
    # Clear cell
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        
    # Vertical center
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Name (Destacado)
    name_para = cell.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(item.name.upper())
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x5a, 0x2d, 0x5a)
    
    # Description
    desc_para = cell.add_paragraph()
    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc_para.add_run(item.description)
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Diet Options
    if item.diet_options:
        diet_para = cell.add_paragraph()
        diet_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = diet_para.add_run(f"({item.diet_options.upper()})")
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Logo (Bottom right)
    if os.path.exists(LOGO_PATH):
        logo_para = cell.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = logo_para.add_run()
        run.add_picture(LOGO_PATH, width=Cm(1.2))

def create_grid_page(doc):
    """Creates a 3x2 table grid for a page."""
    table = doc.add_table(rows=3, cols=2)
    # table.style = 'Table Grid' # Change to None or custom style for no borders if requested
    # We want borders to match the user's design
    
    # Set heights to approx 5cm to 6cm
    for row in table.rows:
        row.height = Cm(5.5)
        
    # Set widths to 8cm
    for col in table.columns:
        col.width = Cm(9) # Slightly wider than 8cm to fill page
        
    # Merge cells in the middle row for the center element
    table.cell(1, 0).merge(table.cell(1, 1))
    
    return table

def generate_individual_signs_docx(request: IndividualSignRequest) -> BytesIO:
    doc = Document()
    
    # Set Page Orientation to Landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Update height/width for landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    # Margins (1cm)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    
    items = []
    for meal in request.meals:
        if meal.menu_name and meal.menu_name.strip():
            items.append(ItemData(
                meal.menu_name.strip(), 
                meal.menu_desc.strip() if meal.menu_desc else "", 
                meal.menu_diet.strip() if meal.menu_diet else ""
            ))
            
        for i in range(1, 11):
            name = getattr(meal, f"menu_{i}_name")
            if name and name.strip():
                desc = getattr(meal, f"menu_{i}_desc")
                diet = getattr(meal, f"menu_{i}_diet")
                items.append(ItemData(
                    name.strip(), 
                    desc.strip() if desc else "", 
                    diet.strip() if diet else ""
                ))
    
    total_items = len(items)
    
    for i in range(0, total_items, 5):
        if i > 0:
            doc.add_page_break()
            
        table = create_grid_page(doc)
        page_items = items[i:i+5]
        
        # Pattern:
        # R0C0 (Item 1), R0C1 (Item 2)
        # R1 (Merged) (Item 3)
        # R2C0 (Item 4), R2C1 (Item 5)
        
        # Mapping indices to table cells
        # cell_map: (page_index) -> (row, col)
        cell_map = {
            0: (0, 0),
            1: (0, 1),
            2: (1, 0), # This is the merged center cell
            3: (2, 0),
            4: (2, 1)
        }
        
        for idx, item in enumerate(page_items):
            row, col = cell_map[idx]
            cell = table.cell(row, col)
            format_cell(cell, item)
            
            # Apply border (if not already in style)
            # In python-docx, borders are tricky without manual XML or a style with borders.
            # I'll rely on the table having some visible border or leave it to the template if we use one.
            # For now, I'll just leave it and see if the user wants borders.
            
    # Save to memory stream
    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    
    return docx_stream
