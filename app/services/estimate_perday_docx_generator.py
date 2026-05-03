import os
import logging
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.schemas.estimate_total import EstimateTotalRequest

logger = logging.getLogger(__name__)

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "estimate_perday_template.docx")

class EstimatePerDayDocxGenerator:
    def __init__(self, template_path=TEMPLATE_PATH):
        self.template_path = template_path
        self.font_name = "Open Sans"
        self.primary_color = 0x612d4b  # Wine color from HTML
        self.text_color = 0x333333     # Main text color
        self.desc_color = 0x555555     # Description color

    def _set_run_font(self, run, size_pt=Pt(10), bold=None, italic=None, color_rgb=None, underline=None):
        rPr = run._element.get_or_add_rPr()
        run.font.name = self.font_name
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), self.font_name)
        rFonts.set(qn('w:hAnsi'), self.font_name)
        rFonts.set(qn('w:cs'), self.font_name)

        if size_pt is not None:
            run.font.size = size_pt
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline
        if color_rgb is not None:
            run.font.color.rgb = RGBColor((color_rgb >> 16) & 0xff, (color_rgb >> 8) & 0xff, color_rgb & 0xff)

    def _format_currency(self, val):
        if val is None:
            return ""
        if isinstance(val, (int, float)):
            s = f"{abs(val):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            if val < 0:
                return f"-$ {s}"
            return f"$ {s}"

        s = str(val).strip()
        if not s:
            return ""
        if not any(c.isdigit() for c in s):
            return s

        is_negative = False
        if s.startswith("-"):
            is_negative = True
            s = s[1:].strip()
            
        s = s.replace("$", "").strip()
        if is_negative:
            return f"-$ {s}"
        return f"$ {s}"

    def _parse_price(self, val):
        if not val:
            return 0.0
        if isinstance(val, (int, float)):
            return float(val)
        clean = str(val).replace("$", "").replace(" ", "").strip()
        if "," in clean and "." in clean:
            if clean.rfind(",") > clean.rfind("."): 
                clean = clean.replace(".", "").replace(",", ".")
            else: 
                clean = clean.replace(",", "")
        elif "," in clean:
            parts = clean.split(",")
            if len(parts[-1]) == 2: 
                clean = clean.replace(",", ".")
            else: 
                clean = clean.replace(",", "")
        try:
            return float(clean)
        except (ValueError, TypeError):
            return 0.0

    def _parse_percentage(self, val):
        if not val:
            return 0.0
        if isinstance(val, (int, float)):
            return float(val) / 100.0
        
        clean = str(val).replace("%", "").strip()
        try:
            return float(clean) / 100.0
        except (ValueError, TypeError):
            return 0.0

    def _replace_placeholders(self, doc, request: EstimateTotalRequest, daily_guests_str: str):
        replacements = {
            "{{EVENT_NAME}}": request.event.name,
            "{{CLIENT_NAME}}": request.client.name,
            "{{CLIENT_ADDRESS}}": request.client.address,
            "{{CLIENT_EMAIL}}": request.client.email,
            "{{REPRESENTATIVE_NAME}}": request.client_representative.name,
            "{{REPRESENTATIVE_EMAIL}}": request.client_representative.email,
            "{{REPRESENTATIVE_PHONE}}": request.client_representative.formatted_phone,
            "{{EVENT_ADDRESS}}": request.event.address,
            "{{EVENT_CODE}}": request.event.code,
            "{{EVENT_START}}": request.event.date_formatted,
            "{{EVENT_END}}": request.event.end_date_formatted,
            "{{DAILY_GUESTS}}": daily_guests_str,
            "{{SERVICE_CHARGE_RATE}}": request.financials.service_charge_rate,
            "{{TAX_NAME}}": request.financials.tax_name,
            "{{TAX_RATE}}": request.financials.tax_rate,
        }

        def process_paragraphs(paragraphs):
            for paragraph in paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        found_in_run = False
                        for run in paragraph.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, str(value or ""))
                                found_in_run = True
                        if not found_in_run and len(paragraph.runs) > 1:
                            full_text = "".join(r.text for r in paragraph.runs)
                            if key in full_text:
                                paragraph.runs[0].text = full_text.replace(key, str(value or ""))
                                for i in range(1, len(paragraph.runs)):
                                    paragraph.runs[i].text = ""

        process_paragraphs(doc.paragraphs)

        def process_tables(tables):
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_paragraphs(cell.paragraphs)
                        if cell.tables:
                            process_tables(cell.tables)

        process_tables(doc.tables)

        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    process_paragraphs(header.paragraphs)
                    process_tables(header.tables)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    process_paragraphs(footer.paragraphs)
                    process_tables(footer.tables)
            
            for part in [section.header, section.first_page_header, section.even_page_header, section.footer, section.first_page_footer, section.even_page_footer]:
                if part:
                    for t in part._element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                        if t.text:
                            for key, value in replacements.items():
                                if key in t.text:
                                    t.text = t.text.replace(key, str(value or ""))

    def generate_docx(self, request: EstimateTotalRequest) -> BytesIO:
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template not found at {self.template_path}")

        doc = Document(self.template_path)

        # De-duplicate meals to avoid repetitions
        unique_meals_map = {} 
        for m in request.meals:
            m_dict = m.model_dump()
            m_dict.pop('show_date_header', None)
            m_dict.pop('show_date_header_2', None)
            m_dict.pop('show_guest_header', None)
            sig = str(sorted(m_dict.items(), key=lambda x: x[0]))
            if sig not in unique_meals_map:
                unique_meals_map[sig] = m
            else:
                if m.show_date_header:
                    unique_meals_map[sig] = m
        
        unique_meals = []
        seen_sigs = set()
        for m in request.meals:
            m_dict = m.model_dump()
            m_dict.pop('show_date_header', None)
            m_dict.pop('show_date_header_2', None)
            m_dict.pop('show_guest_header', None)
            sig = str(sorted(m_dict.items(), key=lambda x: x[0]))
            if sig not in seen_sigs:
                seen_sigs.add(sig)
                unique_meals.append(unique_meals_map[sig])

        # Build DAILY_GUESTS string
        daily_guests_lines = []
        seen_guests_dates = set()
        for meal in unique_meals:
            if meal.show_date_header_2 and meal.date_day_name:
                key = (meal.date_day_name, meal.guest_count)
                if key not in seen_guests_dates:
                    seen_guests_dates.add(key)
                    daily_guests_lines.append(f"{meal.date_day_name} {meal.guest_count} Guests")
        daily_guests_str = "\n".join(daily_guests_lines)

        self._replace_placeholders(doc, request, daily_guests_str)

        marker_para = None
        for p in doc.paragraphs:
            if "[DYNAMIC_CONTENT_START]" in p.text:
                marker_para = p
                break

        if not marker_para:
            logger.warning("Marker [DYNAMIC_CONTENT_START] not found. Appending to end.")

        def add_p(text="", alignment=None, space_after=Pt(6), space_before=Pt(0), bold=False, italic=False, size=Pt(10), color=0x333333, underline=False):
            if marker_para:
                p = marker_para.insert_paragraph_before(text)
            else:
                p = doc.add_paragraph(text)
            
            if alignment: p.alignment = alignment
            p.paragraph_format.space_after = space_after
            p.paragraph_format.space_before = space_before
            
            if p.runs:
                run = p.runs[0]
            else:
                run = p.add_run(text)
            
            if not p.runs and text:
                run = p.add_run(text)

            self._set_run_font(run, size_pt=size, bold=bold, italic=italic, color_rgb=color, underline=underline)
            return p

        def add_hr():
            p = add_p(space_after=Pt(3), space_before=Pt(0), size=Pt(1))
            p.paragraph_format.line_spacing = Pt(1)
            p_pr = p._element.get_or_add_pPr()
            p_bdr = p_pr.find(qn('w:pBdr'))
            if p_bdr is None:
                p_bdr = OxmlElement('w:pBdr')
                p_pr.insert(0, p_bdr)
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            p_bdr.append(bottom)

        # --- MENU SECTION ---
        add_p("MENUS", bold=True, size=Pt(10), color=self.primary_color, space_after=Pt(0))
        add_p(request.event.date_formatted, space_after=Pt(0))

        if request.event.dietary_restrictions:
            add_p("Dietary Restrictions", bold=True, size=Pt(10), color=self.primary_color, space_after=Pt(0))
            add_p(request.event.dietary_restrictions, space_after=Pt(0))

        for meal in unique_meals:
            if meal.show_date_header:
                add_p(meal.date_header, bold=True, space_before=Pt(6))
                add_hr()
            
            cat_text = meal.category_name.upper()
            if meal.time_range:
                cat_text += f": {meal.time_range}"
            
            add_p(cat_text, bold=True, size=Pt(10), color=self.primary_color, space_before=Pt(8))
            
            if meal.provide_by_client:
                p = add_p("◽ Provided by client", space_before=Pt(4))
                p.paragraph_format.left_indent = Cm(0.5)
                continue

            if meal.description:
                add_p(meal.description, italic=True, space_before=Pt(0))

            grouped_subs = [] 
            for i in range(1, 13):
                s_name = getattr(meal, f"subcategory_{i}_name", "").strip()
                s_desc = getattr(meal, f"subcategory_{i}_description", "").strip()
                s_items = getattr(meal, f"subcategory_{i}_items", [])

                if not s_name and not s_items:
                    continue

                found = False
                if s_name:
                    for gs in grouped_subs:
                        if gs['name'] == s_name:
                            existing_names = {it.name for it in gs['items']}
                            for it in s_items:
                                if it.name not in existing_names:
                                    gs['items'].append(it)
                                    existing_names.add(it.name)
                            if s_desc and s_desc not in gs['desc']:
                                gs['desc'] = (gs['desc'] + " " + s_desc).strip()
                            found = True
                            break
                
                if not found:
                    grouped_subs.append({
                        'name': s_name,
                        'desc': s_desc,
                        'items': list(s_items)
                    })

            for gs in grouped_subs:
                sub_name = gs['name']
                sub_desc = gs['desc']
                sub_items = gs['items']

                if sub_name:
                    sub_p = add_p(sub_name, bold=True, space_before=Pt(6))
                    sub_p.runs[0].underline = True

                if sub_desc:
                    add_p(sub_desc, size=Pt(10), italic=True)

                for item in sub_items:
                    item_p = add_p(space_after=Pt(2))
                    item_p.paragraph_format.left_indent = Cm(0.8)
                    item_p.paragraph_format.first_line_indent = Cm(-0.4)
                    
                    r_bullet = item_p.add_run("◽ ")
                    self._set_run_font(r_bullet, bold=True)
                    
                    r_name = item_p.add_run(item.name)
                    self._set_run_font(r_name, bold=False, underline=False)

                    if item.description:
                        desc_p = add_p(item.description, size=Pt(10), italic=True, color=self.desc_color, space_after=Pt(4))
                        desc_p.paragraph_format.left_indent = Cm(1.2)

        # --- FINANCIAL SECTION ---
        add_p(space_before=Pt(10))
        add_p("PROPOSAL OF SERVICES", bold=True, size=Pt(10), color=self.primary_color, space_after=Pt(0), space_before=Pt(10))
        add_p(request.event.end_date_formatted, space_after=Pt(0)) 

        # 1. Food Service
        add_p("Food Service", bold=True, size=Pt(10), color=self.primary_color, space_after=Pt(0), space_before=Pt(10))
        
        daily_food_totals = {}
        for m in unique_meals:
            if not m.provide_by_client:
                val = self._parse_price(m.total_category_precio_guest_por_dia)
                daily_food_totals[m.date_header] = daily_food_totals.get(m.date_header, 0.0) + val

        for meal in unique_meals:
            if meal.show_date_header:
                add_p(meal.date_header, bold=True, space_before=Pt(6), space_after=Pt(0))
                add_hr()
                
                total_val = daily_food_totals.get(meal.date_header, 0.0)
                if total_val >= 0:
                    add_p(self._format_currency(total_val), bold=True, space_after=Pt(4))
                    
            if meal.show_guest_header:
                add_p(f"Based on {meal.guest_count} Guests", size=Pt(10), italic=True, space_before=Pt(0))
            
            p = add_p(space_after=Pt(2))
            p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
            r_label = p.add_run(meal.category_precio_guest)
            self._set_run_font(r_label)
            r_spacer = p.add_run("\t") 
            self._set_run_font(r_spacer)
            if not meal.provide_by_client:
                r_val = p.add_run(self._format_currency(meal.total_category_precio_guest_por_dia))
                self._set_run_font(r_val, bold=True)
            else:
                r_client = p.add_run("Provided by client")
                self._set_run_font(r_client, italic=True)

        # 2. Labor
        if request.labor_services:
            add_p("Labor Service Fees", bold=True, size=Pt(10), color=self.primary_color, space_before=Pt(15), space_after=Pt(0))
            
            seen_labor = set()
            unique_labor = []
            for labor in request.labor_services:
                key = (labor.date_header, labor.hours, labor.name, labor.total)
                if key not in seen_labor:
                    seen_labor.add(key)
                    unique_labor.append(labor)

            labor_groups = []
            for labor in unique_labor:
                found = False
                for g in labor_groups:
                    if g['date'] == labor.date_header and g['hours'] == labor.hours:
                        g['items'].append(labor)
                        found = True
                        break
                if not found:
                    labor_groups.append({
                        'date': labor.date_header,
                        'hours': labor.hours,
                        'show_date': labor.show_date_header,
                        'items': [labor]
                    })

            for group in labor_groups:
                if group['show_date']:
                    add_p(group['date'], bold=True, space_before=Pt(6), space_after=Pt(0))
                    add_hr()
                
                p_desc = add_p(space_after=Pt(0))
                r_header_prefix = p_desc.add_run("Staff suggested based on ")
                self._set_run_font(r_header_prefix, italic=True)
                
                r_hours = p_desc.add_run(f"{group['hours']}")
                self._set_run_font(r_hours, italic=True, bold=True)
                
                r_header_suffix = p_desc.add_run(" hours of labor. ")
                self._set_run_font(r_header_suffix, italic=True)
                
                names_str = ", ".join([item.name for item in group['items']])
                r_names = p_desc.add_run(names_str)
                self._set_run_font(r_names, bold=False)
                
                total_val = sum(self._parse_price(item.total) for item in group['items'])
                p_total = add_p(self._format_currency(total_val), alignment=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, space_after=Pt(4))
                self._set_run_font(p_total.runs[0], bold=True, color_rgb=0x000000)

        # 3. Extras
        if request.extras_events:
            add_p("Extras Services", bold=True, size=Pt(10), color=self.primary_color, space_before=Pt(15), space_after=Pt(0))
            
            seen_extras = set()
            unique_extras = []
            for extra in request.extras_events:
                key = (extra.date_header, extra.is_rental, extra.is_sales, extra.name, extra.name_rental, extra.name_sales, extra.total, extra.provide_by_client)
                if key not in seen_extras:
                    seen_extras.add(key)
                    unique_extras.append(extra)

            for extra in unique_extras:
                if extra.show_date_header:
                    add_p(extra.date_header, bold=True, space_after=Pt(0))
                    add_hr()
                
                if extra.is_rental:
                    add_p("Rentals", bold=True, space_before=Pt(6))
                
                if extra.is_sales:
                    add_p("Sales", bold=True, space_before=Pt(6))

                p = add_p(space_after=Pt(2))
                p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
                
                display_name = extra.name
                if not extra.provide_by_client:
                    if extra.is_rental: display_name = extra.name_rental
                    elif extra.is_sales: display_name = extra.name_sales

                if extra.provide_by_client:
                    txt = f"{display_name}\tProvide by the client"
                else:
                    txt = f"{display_name}\t{self._format_currency(extra.total)}"
                p_extra = p.add_run(txt)
                self._set_run_font(p_extra, bold=True)

        # 4. Final Summary
        add_p("Cost of Balance", bold=True, size=Pt(10), color=self.primary_color, space_before=Pt(10))
        fin = request.financials
        
        # --- RE-CALCULATION ENGINE (Excel Model) ---
        # 1. Base Components
        real_food_total = sum(daily_food_totals.values())
        
        real_labor_total = 0.0
        if request.labor_services:
            for group in labor_groups:
                real_labor_total += sum(self._parse_price(item.total) for item in group['items'])
            
        real_extras_sales_total = 0.0
        real_extras_rentals_total = 0.0
        if request.extras_events:
            for ex in unique_extras:
                if not ex.provide_by_client:
                    val = self._parse_price(ex.total)
                    if ex.is_sales: real_extras_sales_total += val
                    if ex.is_rental: real_extras_rentals_total += val

        real_gratuity = self._parse_price(fin.gratuity)
        real_discount = self._parse_price(fin.discount)
        real_donation = self._parse_price(fin.donation)

        # 2. SubTotal 1 & Taxes
        # SubTotal 1 = Food + Labor + Extras Sales + Gratuity - Discount - Donation
        subtotal_1 = (real_food_total + real_labor_total + real_extras_sales_total + 
                      real_gratuity - abs(real_discount) - abs(real_donation))
        
        tax_rate = self._parse_percentage(fin.tax_rate)
        real_tax = subtotal_1 * tax_rate
        subtotal_2 = subtotal_1 + real_tax

        # 3. Service Charge & SubTotal 4
        # Service Charge = (Food + Labor) * Rate
        service_charge_rate = self._parse_percentage(fin.service_charge_rate)
        real_service_charge = (real_food_total + real_labor_total) * service_charge_rate
        
        # Subtotal 4 = Subtotal 2 + Service Charge (Rentals are informative only)
        subtotal_4 = subtotal_2 + real_service_charge

        # 4. Credit Card & Final Total
        cc_rate = self._parse_percentage(fin.credit_card_percent)
        real_cc_fee = subtotal_4 * cc_rate
        real_grand_total = subtotal_4 + real_cc_fee

        # --- RENDER SUMMARY ---
        summary_items = [
            ("Food", real_food_total, True),
            ("Labor Cost", real_labor_total, True),
            ("Extras Services (Sales)", real_extras_sales_total, True),
            ("Gratuity", real_gratuity, False),
            ("Discount", -abs(real_discount), False),
            ("Donation", -abs(real_donation), False),
            ("SubTotal 1", subtotal_1, True, True), # Label, value, always, is_bold
            (f"{fin.tax_rate} {fin.tax_name}", real_tax, True),
            ("SubTotal 2", subtotal_2, True, True),
            ("Extras Services (Rentals)", real_extras_rentals_total, True, False, True), # informative only
            (f"{fin.service_charge_rate} Service Charge", real_service_charge, True),
            ("SubTotal 4", subtotal_4, True, True),
            ("Credit Card Fee", real_cc_fee, False),
        ]

        for item in summary_items:
            label = item[0]
            val = item[1]
            show_always = item[2]
            is_bold_item = item[3] if len(item) > 3 else False
            is_informative = item[4] if len(item) > 4 else False

            if not show_always and abs(val) < 0.01:
                continue
                
            p = add_p(space_after=Pt(2))
            p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
            
            r_label = p.add_run(label)
            self._set_run_font(r_label, bold=is_bold_item)
            
            r_tab = p.add_run("\t")
            self._set_run_font(r_tab)
            
            formatted_val = self._format_currency(val)
            if is_informative:
                formatted_val = f"({formatted_val}*) "
            
            r_val = p.add_run(formatted_val)
            self._set_run_font(r_val, bold=is_bold_item)

        # Final Total Line
        total_p = add_p(space_after=Pt(2), space_before=Pt(8))
        total_p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
        
        r_total = total_p.add_run(f"Final\t{self._format_currency(real_grand_total)}")
        self._set_run_font(r_total, bold=True, size_pt=Pt(10), color_rgb=self.primary_color)
        
        p_pr = total_p._element.get_or_add_pPr()
        p_bdr = p_pr.find(qn('w:pBdr'))
        if p_bdr is None:
            p_bdr = OxmlElement('w:pBdr')
            p_pr.insert(0, p_bdr)
        
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12')
        top.set(qn('w:color'), '612D4B')
        p_bdr.append(top)

        if marker_para:
            p_element = marker_para._element
            p_element.getparent().remove(p_element)

        docx_stream = BytesIO()
        doc.save(docx_stream)
        docx_stream.seek(0)
        return docx_stream
